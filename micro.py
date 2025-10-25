from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Word document
doc = Document()

# Add a title
doc.add_heading('DEMONSTRATION OF ELECTROMAGNET AND ITS APPLICATIONS', 0)

# Add a cover info paragraph
doc.add_paragraph('Micro Project Report\n\nSubmitted by: __________________________\nInstitution: __________________________\nClass: __________________________')

# Add a page break
doc.add_page_break()

# Add sections with headings and content
sections = {
    'Introduction': 'An electromagnet is a temporary magnet formed when an electric current passes through a coil of wire wrapped around a soft iron core.',
    'Objective': 'To demonstrate the principles of electromagnetic induction, and see how current and coil turns affect strength.',
    'Materials Required': 'Iron nail, insulated copper wire, battery, switch, metallic clips.',
    'Theory': 'When current flows through a coil, a magnetic field is produced. The strength of this field depends on the current and number of turns.',
    'Construction/Method': '1. Wrap wire tightly around nail.\n2. Connect ends to battery.\n3. Observe magnetic attraction to clips.\n4. Repeat with more turns.',
    'Observation': 'Increase in wire turns or current results in a stronger attraction. Record results as a data table.',
    'Applications': 'Motors, relays, bells, cranes, MRI, magnetic locks.',
    'Advantages & Disadvantages': 'Advantages: controllable magnetism, switchable, efficient.\nDisadvantages: needs power, heats with use.',
    'Precautions': 'Disconnect after use. Do not leave powered for long.\nEnsure insulation and correct connections.',
    'Conclusion': 'Magnetic field strength depends on current and number of wire turns. Electromagnets are widely useful.',
    'References': '1. TeachEngineering.org\n2. GeeksforGeeks.org\n3. BYJU’S\n4. Vedantu\n5. UMD Physics Outreach'
}

for title, body in sections.items():
    doc.add_heading(title, level=1)
    doc.add_paragraph(body)
    doc.add_page_break()

# Insert an image (diagram.png must exist in your directory!)
doc.add_heading('Electromagnet Diagram', level=2)
doc.add_picture('diagram.png', width=Inches(4))
doc.add_paragraph('Figure: Labeled diagram of basic electromagnet setup.')

# Save the document
doc.save('Electromagnet_Project_Report.docx')
print("Word file created as Electromagnet_Project_Report.docx")
